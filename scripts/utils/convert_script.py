"""Convert presentation script to Word/PDF"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

script_dir = Path(__file__).parent
md_file = script_dir / 'VIDEO_PRESENTATION_SCRIPT.md'

with open(md_file, 'r', encoding='utf-8') as f:
    content = f.read()

doc = Document()

# Title
title = doc.add_paragraph()
title_run = title.add_run('ShiftOps - Video Presentation Script')
title_run.bold = True
title_run.font.size = Pt(28)
title_run.font.color.rgb = RGBColor(30, 60, 114)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
sub_run = subtitle.add_run('Leadership & Management Demo Guide')
sub_run.font.size = Pt(16)
sub_run.font.color.rgb = RGBColor(100, 100, 100)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Process content
lines = content.split('\n')
for line in lines:
    if line.startswith('# '):
        continue  # Skip main title
    elif line.startswith('## '):
        h = doc.add_heading(line[3:], level=1)
    elif line.startswith('### '):
        h = doc.add_heading(line[4:], level=2)
    elif line.startswith('> '):
        p = doc.add_paragraph()
        r = p.add_run(line[2:])
        r.italic = True
        r.font.color.rgb = RGBColor(0, 100, 0)
    elif line.startswith('- [ ]') or line.startswith('- [x]'):
        p = doc.add_paragraph(line[6:], style='List Bullet')
    elif line.startswith('- '):
        p = doc.add_paragraph(line[2:], style='List Bullet')
    elif line.startswith('| ') and '---' not in line:
        p = doc.add_paragraph(line)
        if p.runs:
            p.runs[0].font.size = Pt(10)
    elif line.strip().startswith('"'):
        p = doc.add_paragraph()
        r = p.add_run(line.strip())
        r.italic = True
    elif line.strip():
        p = doc.add_paragraph(line)

output_docx = script_dir / 'ShiftOps_Presentation_Script.docx'
doc.save(output_docx)
print(f'[OK] Created: {output_docx}')

try:
    from docx2pdf import convert
    output_pdf = script_dir / 'ShiftOps_Presentation_Script.pdf'
    convert(str(output_docx), str(output_pdf))
    print(f'[OK] Created: {output_pdf}')
except Exception as e:
    print(f'[WARNING] PDF conversion failed: {e}')




