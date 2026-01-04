"""
Markdown to Word/PDF Converter
Converts USER_GUIDE.md and ADMIN_GUIDE.md to Word and PDF documents
"""

import os
import re
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[WARNING] python-docx not installed. Run: pip install python-docx")

try:
    from docx2pdf import convert
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("[WARNING] docx2pdf not installed. Run: pip install docx2pdf")


class MarkdownToDocxConverter:
    def __init__(self, screenshots_dir):
        self.screenshots_dir = Path(screenshots_dir)
        self.doc = None
        self.current_list_level = 0
        
    def create_document(self):
        """Initialize a new Word document with styles."""
        self.doc = Document()
        
        # Set default font
        style = self.doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # Configure heading styles
        for i in range(1, 5):
            heading_style = self.doc.styles[f'Heading {i}']
            heading_style.font.name = 'Calibri'
            heading_style.font.color.rgb = RGBColor(30, 60, 114)
            heading_style.font.bold = True
            if i == 1:
                heading_style.font.size = Pt(24)
            elif i == 2:
                heading_style.font.size = Pt(18)
            elif i == 3:
                heading_style.font.size = Pt(14)
            else:
                heading_style.font.size = Pt(12)
        
        return self.doc
    
    def add_title_page(self, title, subtitle="Shift Handover Application"):
        """Add a professional title page."""
        # Add spacing at top
        for _ in range(4):
            self.doc.add_paragraph()
        
        # Main title
        title_para = self.doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(36)
        title_run.font.color.rgb = RGBColor(30, 60, 114)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        self.doc.add_paragraph()
        subtitle_para = self.doc.add_paragraph()
        subtitle_run = subtitle_para.add_run(subtitle)
        subtitle_run.font.size = Pt(18)
        subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing
        for _ in range(4):
            self.doc.add_paragraph()
        
        # Version and date
        info_para = self.doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_run = info_para.add_run(f"Version: 2.0.0\nGenerated: {datetime.now().strftime('%B %d, %Y')}")
        info_run.font.size = Pt(12)
        info_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Company
        self.doc.add_paragraph()
        company_para = self.doc.add_paragraph()
        company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        company_run = company_para.add_run("EPAM Systems")
        company_run.font.size = Pt(14)
        company_run.font.color.rgb = RGBColor(30, 60, 114)
        company_run.bold = True
        
        self.doc.add_page_break()
    
    def parse_markdown(self, md_content):
        """Parse markdown content and convert to Word."""
        lines = md_content.split('\n')
        i = 0
        in_code_block = False
        code_content = []
        in_table = False
        table_rows = []
        
        while i < len(lines):
            line = lines[i]
            
            # Handle code blocks
            if line.strip().startswith('```'):
                if in_code_block:
                    # End code block
                    self.add_code_block('\n'.join(code_content))
                    code_content = []
                    in_code_block = False
                else:
                    # Start code block
                    in_code_block = True
                i += 1
                continue
            
            if in_code_block:
                code_content.append(line)
                i += 1
                continue
            
            # Handle tables
            if '|' in line and not line.strip().startswith('<!--'):
                if not in_table:
                    in_table = True
                    table_rows = []
                
                # Skip separator lines (|---|---|)
                if re.match(r'^[\|\-\:\s]+$', line.strip()):
                    i += 1
                    continue
                
                # Parse table row
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                if cells:
                    table_rows.append(cells)
                i += 1
                continue
            elif in_table:
                # End of table
                if table_rows:
                    self.add_table(table_rows)
                in_table = False
                table_rows = []
            
            # Skip HTML comments
            if line.strip().startswith('<!--'):
                i += 1
                continue
            
            # Handle horizontal rules
            if line.strip() in ['---', '***', '___']:
                self.add_horizontal_rule()
                i += 1
                continue
            
            # Handle headings
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2)
                # Remove markdown links from heading
                text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
                self.add_heading(text, level)
                i += 1
                continue
            
            # Handle images
            img_match = re.match(r'!\[([^\]]*)\]\(([^\)]+)\)', line.strip())
            if img_match:
                alt_text = img_match.group(1)
                img_path = img_match.group(2)
                self.add_image(img_path, alt_text)
                i += 1
                continue
            
            # Handle blockquotes (> text)
            if line.strip().startswith('>'):
                quote_text = line.strip()[1:].strip()
                self.add_blockquote(quote_text)
                i += 1
                continue
            
            # Handle unordered lists
            list_match = re.match(r'^(\s*)[-*]\s+(.+)$', line)
            if list_match:
                indent = len(list_match.group(1))
                text = list_match.group(2)
                self.add_list_item(text, indent // 2)
                i += 1
                continue
            
            # Handle ordered lists
            ordered_match = re.match(r'^(\s*)(\d+)\.\s+(.+)$', line)
            if ordered_match:
                indent = len(ordered_match.group(1))
                text = ordered_match.group(3)
                self.add_list_item(text, indent // 2, ordered=True)
                i += 1
                continue
            
            # Handle regular paragraphs
            if line.strip():
                self.add_paragraph(line)
            
            i += 1
        
        # Handle any remaining table
        if in_table and table_rows:
            self.add_table(table_rows)
    
    def add_heading(self, text, level):
        """Add a heading to the document."""
        level = min(level, 4)  # Max heading level in Word
        heading = self.doc.add_heading(text, level=level)
        
        # Add page break before major sections
        if level == 1 and len(self.doc.paragraphs) > 5:
            heading.paragraph_format.page_break_before = True
    
    def add_paragraph(self, text):
        """Add a paragraph with inline formatting."""
        para = self.doc.add_paragraph()
        
        # Process inline formatting
        self.process_inline_formatting(para, text)
        
        para.paragraph_format.space_after = Pt(6)
    
    def process_inline_formatting(self, para, text):
        """Process bold, italic, code, and links in text."""
        # Pattern to match various inline elements
        pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^\)]+\)|[^*`\[]+)'
        
        parts = re.findall(pattern, text)
        
        for part in parts:
            if not part:
                continue
            
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                run = para.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                # Italic text
                run = para.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                # Inline code
                run = para.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(128, 0, 0)
            elif part.startswith('['):
                # Link - extract text only
                link_match = re.match(r'\[([^\]]+)\]\([^\)]+\)', part)
                if link_match:
                    run = para.add_run(link_match.group(1))
                    run.font.color.rgb = RGBColor(0, 102, 204)
                    run.underline = True
            else:
                # Regular text
                para.add_run(part)
    
    def add_list_item(self, text, level=0, ordered=False):
        """Add a list item."""
        para = self.doc.add_paragraph()
        para.style = 'List Bullet' if not ordered else 'List Number'
        
        # Set indentation based on level
        para.paragraph_format.left_indent = Inches(0.25 * (level + 1))
        
        self.process_inline_formatting(para, text)
    
    def add_blockquote(self, text):
        """Add a blockquote (styled paragraph)."""
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.5)
        
        # Add a vertical bar or styling
        run = para.add_run(text)
        run.italic = True
        run.font.color.rgb = RGBColor(80, 80, 80)
        
        # Add light background shading
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'E8F4F8')
        para._p.get_or_add_pPr().append(shading)
    
    def add_code_block(self, code):
        """Add a code block."""
        para = self.doc.add_paragraph()
        
        # Style the code block
        run = para.add_run(code)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        
        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        
        # Add shading
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'F5F5F5')
        para._p.get_or_add_pPr().append(shading)
    
    def add_table(self, rows):
        """Add a table to the document."""
        if not rows:
            return
        
        num_cols = max(len(row) for row in rows)
        table = self.doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Table Grid'
        
        for i, row_data in enumerate(rows):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                if j < num_cols:
                    cell = row.cells[j]
                    cell.text = cell_text
                    
                    # Style header row
                    if i == 0:
                        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(cell_text)
                        run.bold = True
                        
                        # Add shading to header
                        shading = OxmlElement('w:shd')
                        shading.set(qn('w:fill'), '1E3C72')
                        cell._tc.get_or_add_tcPr().append(shading)
                        run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Add spacing after table
        self.doc.add_paragraph()
    
    def add_image(self, img_path, alt_text=""):
        """Add an image to the document."""
        # Handle relative paths
        if img_path.startswith('screenshots/'):
            full_path = self.screenshots_dir.parent / img_path
        else:
            full_path = self.screenshots_dir / img_path
        
        if full_path.exists():
            try:
                # Add the image
                self.doc.add_picture(str(full_path), width=Inches(6))
                
                # Center the image
                last_para = self.doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add caption if alt_text provided
                if alt_text:
                    caption = self.doc.add_paragraph()
                    caption_run = caption.add_run(f"Figure: {alt_text}")
                    caption_run.font.size = Pt(9)
                    caption_run.font.italic = True
                    caption_run.font.color.rgb = RGBColor(100, 100, 100)
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                self.doc.add_paragraph()  # Add spacing
            except Exception as e:
                # If image fails, add placeholder
                para = self.doc.add_paragraph()
                run = para.add_run(f"[Image: {alt_text or img_path}]")
                run.font.color.rgb = RGBColor(150, 150, 150)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # Image not found - add placeholder
            para = self.doc.add_paragraph()
            run = para.add_run(f"[Image: {alt_text or img_path}]")
            run.font.color.rgb = RGBColor(150, 150, 150)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def add_horizontal_rule(self):
        """Add a horizontal rule."""
        para = self.doc.add_paragraph()
        para.add_run('_' * 80)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.runs[0].font.color.rgb = RGBColor(200, 200, 200)
    
    def save(self, output_path):
        """Save the document."""
        self.doc.save(output_path)


def convert_md_to_docx(md_file, output_file, screenshots_dir, title):
    """Convert a Markdown file to Word document."""
    print(f"  Reading {md_file}...")
    
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    converter = MarkdownToDocxConverter(screenshots_dir)
    converter.create_document()
    
    # Add title page
    converter.add_title_page(title)
    
    # Parse and convert markdown
    print(f"  Converting content...")
    converter.parse_markdown(md_content)
    
    # Save document
    converter.save(output_file)
    print(f"  [OK] Created: {output_file}")
    
    return True


def convert_to_pdf(docx_file, pdf_file):
    """Convert Word document to PDF."""
    if not PDF_AVAILABLE:
        print("[WARNING] Cannot convert to PDF - docx2pdf not installed")
        return False
    
    try:
        convert(docx_file, pdf_file)
        print(f"  [OK] Created: {pdf_file}")
        return True
    except Exception as e:
        print(f"  [ERROR] PDF conversion failed: {e}")
        return False


def main():
    print("=" * 70)
    print("  Markdown to Word/PDF Converter for ShiftOps Documentation")
    print("=" * 70)
    print()
    
    # Get paths
    script_dir = Path(__file__).parent
    screenshots_dir = script_dir / "screenshots"
    
    user_guide_md = script_dir / "USER_GUIDE.md"
    admin_guide_md = script_dir / "ADMIN_GUIDE.md"
    
    # Check for source files
    if not user_guide_md.exists():
        print(f"[ERROR] User Guide not found: {user_guide_md}")
        return
    
    if not admin_guide_md.exists():
        print(f"[ERROR] Admin Guide not found: {admin_guide_md}")
        return
    
    print(f"[INFO] Screenshots directory: {screenshots_dir}")
    print(f"[INFO] Found {len(list(screenshots_dir.glob('*.png')))} screenshots")
    print()
    
    # Convert User Guide
    print("[1/4] Converting User Guide to Word...")
    user_docx = script_dir / "ShiftOps_User_Guide.docx"
    if convert_md_to_docx(user_guide_md, user_docx, screenshots_dir, "User Guide"):
        print("[2/4] Converting User Guide to PDF...")
        user_pdf = script_dir / "ShiftOps_User_Guide.pdf"
        convert_to_pdf(user_docx, user_pdf)
    
    print()
    
    # Convert Admin Guide
    print("[3/4] Converting Admin Guide to Word...")
    admin_docx = script_dir / "ShiftOps_Admin_Guide.docx"
    if convert_md_to_docx(admin_guide_md, admin_docx, screenshots_dir, "Administrator Guide"):
        print("[4/4] Converting Admin Guide to PDF...")
        admin_pdf = script_dir / "ShiftOps_Admin_Guide.pdf"
        convert_to_pdf(admin_docx, admin_pdf)
    
    print()
    print("=" * 70)
    print("  Documentation conversion complete!")
    print()
    print("  Generated files:")
    print(f"    - ShiftOps_User_Guide.docx")
    print(f"    - ShiftOps_User_Guide.pdf")
    print(f"    - ShiftOps_Admin_Guide.docx")
    print(f"    - ShiftOps_Admin_Guide.pdf")
    print()
    print(f"  Location: {script_dir}")
    print("=" * 70)


if __name__ == "__main__":
    main()


