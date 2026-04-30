"""
Documentation Generator Script
Generates Word and PDF documents with screenshots from the docs/screenshots folder
"""

import os
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[WARNING] python-docx not installed. Run: pip install python-docx")

try:
    from docx2pdf import convert
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("[WARNING] docx2pdf not installed. Run: pip install docx2pdf")

# Screenshot categories and their files
USER_GUIDE_SECTIONS = {
    "Authentication": {
        "description": "Login and authentication screenshots",
        "files": [
            ("login-sso.png", "SSO Login Page", "Default login view with SSO button for single sign-on authentication"),
            ("login-credentials.png", "Credentials Login", "Alternative login using username and password"),
            ("epam-sso.png", "EPAM SSO Portal", "EPAM Single Sign-On authentication page"),
        ]
    },
    "Dashboard": {
        "description": "Main application dashboard",
        "files": [
            ("dashboard.png", "Main Dashboard", "Overview of shifts, incidents, and key metrics"),
        ]
    },
    "Multi-Team Access": {
        "description": "Access multiple teams from a single account",
        "files": [
            ("multi-team-access.png", "Multi-Team Overview", "Overview of multi-team access feature"),
            ("multi-team-reports.png", "Reports with Team Filter", "Filtering reports by team"),
            ("multi-team-config.png", "Multi-Team Configuration", "How team access is configured"),
            ("team-switcher.png", "Team Switcher", "Dropdown to switch between teams"),
        ]
    },
    "Handover Form": {
        "description": "Submit shift handover reports",
        "files": [
            ("handover-form.png", "Handover Form", "Complete handover submission form"),
            ("handover-shift-info.png", "Shift Information", "Date, shift type, and engineer selection"),
            ("handover-incidents.png", "Incidents Section", "Recording incidents during the shift"),
            ("handover-keypoints.png", "Key Points Section", "Important points to highlight"),
            ("handover-changeinfo.png", "Change Info Section", "Recording change implementations"),
            ("handover-kb.png", "KB Updates Section", "Knowledge base updates"),
            ("handover-notes.png", "Additional Notes", "Free-form notes section"),
        ]
    },
    "Reports": {
        "description": "View and export handover reports",
        "files": [
            ("handover-reports.png", "Reports List", "List of all submitted reports"),
            ("reports-filters.png", "Report Filters", "Filter reports by date, team, status"),
            ("report-details.png", "Report Details", "Expanded view of a single report"),
            ("export-buttons.png", "Export Options", "Export to CSV, Excel, or PDF"),
            ("delete-draft.png", "Delete Draft", "Delete draft reports"),
        ]
    },
    "Key Points": {
        "description": "Track ongoing key points and updates",
        "files": [
            ("keypoints.png", "Key Points Page", "Overview of all key points"),
            ("keypoints-list.png", "Key Points List", "Cards showing status badges"),
            ("keypoint-update.png", "Add Update", "Form to add updates to key points"),
        ]
    },
    "Tools": {
        "description": "Additional tools for managing data",
        "files": [
            ("change-info.png", "Change Info Page", "Manage change implementations"),
            ("add-change.png", "Add Change Form", "Form to add new change records"),
            ("kb-updates.png", "KB Updates Page", "Knowledge base updates management"),
            ("vendor-details.png", "Vendor Details", "Vendor contact information"),
        ]
    },
    "Roster & Scheduling": {
        "description": "View shift schedules and rosters",
        "files": [
            ("shift-roster.png", "Shift Roster", "Team shift roster view"),
            ("roster-calendar.png", "Calendar View", "Monthly calendar with shifts"),
            ("shift-allowance.png", "Shift Allowance", "Generate allowance reports"),
            ("team-details.png", "Team Details", "Team member information"),
        ]
    },
    "CTask Assignment": {
        "description": "ServiceNow CTask assignments",
        "files": [
            ("ctask-assignment.png", "CTask Page", "CTask assignment overview"),
            ("ctask-list.png", "CTask List", "List of change tasks"),
        ]
    },
    "Support Tools": {
        "description": "On-call and escalation information",
        "files": [
            ("oncall-dashboard.png", "OnCall Dashboard", "On-call engineer information"),
            ("escalation-matrix.png", "Escalation Matrix", "Escalation contact information"),
            ("servicenow-integration.png", "ServiceNow Integration", "ServiceNow connectivity"),
        ]
    },
    "User Management": {
        "description": "User administration",
        "files": [
            ("user-management.png", "User Management", "Manage users in your team"),
            ("audit-logs.png", "Audit Logs", "Activity and audit trail"),
        ]
    },
    "Account Settings": {
        "description": "Personal account settings",
        "files": [
            ("my-profile.png", "My Profile", "View and edit your profile"),
            ("change-password.png", "Change Password", "Update your password"),
            ("account-settings.png", "Account Settings", "Personal preferences"),
            ("notifications.png", "Notifications", "View notifications"),
            ("system-alerts.png", "System Alerts", "System status and alerts"),
        ]
    },
    "Help & About": {
        "description": "Help documentation and application information",
        "files": [
            ("help-support.png", "Help & Support", "Get help and support"),
            ("about.png", "About", "Application version and information"),
        ]
    },
}

ADMIN_GUIDE_SECTIONS = {
    "Admin Role Dashboards": {
        "description": "Different dashboard views based on admin role",
        "files": [
            ("admin-super-dashboard.png", "Super Admin Dashboard", "Full system access dashboard"),
            ("admin-account-dashboard.png", "Account Admin Dashboard", "Account-level management view"),
            ("admin-team-dashboard.png", "Team Admin Dashboard", "Team-level management view"),
            ("admin-role-comparison.png", "Role Comparison", "Permissions comparison table"),
        ]
    },
    "Account Management": {
        "description": "Super Admin - Manage accounts",
        "files": [
            ("admin-accounts.png", "Accounts List", "All accounts in the system"),
            ("admin-create-account.png", "Create Account", "New account creation form"),
        ]
    },
    "User Management": {
        "description": "Manage users across the organization",
        "files": [
            ("admin-users-global.png", "Global Users", "All users (Super Admin view)"),
            ("admin-users-account.png", "Account Users", "Users in specific account"),
            ("admin-create-user.png", "Create User", "New user creation form"),
            ("admin-user-edit.png", "Edit User", "Modify user details"),
            ("admin-role-assign.png", "Role Assignment", "Assign user roles"),
            ("admin-user-deactivate.png", "Deactivate User", "Disable user access"),
            ("admin-password-reset.png", "Password Reset", "Reset user password"),
            ("admin-user-create-workflow.png", "User Creation Workflow", "Step-by-step user creation"),
        ]
    },
    "Team Management": {
        "description": "Manage teams and team members",
        "files": [
            ("admin-teams-account.png", "Teams List", "Teams in account"),
            ("admin-team-config-options.png", "Team Configuration", "Team settings options"),
            ("admin-user-linking.png", "User Linking", "Link users to team members"),
            ("admin-member-linking.png", "Member Linking", "Link team members to users"),
        ]
    },
    "System Configuration": {
        "description": "Super Admin - System-wide configuration",
        "files": [
            ("admin-config-menu.png", "Configuration Menu", "Configuration navigation"),
            ("admin-email-config.png", "Email Configuration", "SMTP settings"),
            ("admin-email-test.png", "Email Test", "Test email functionality"),
            ("admin-team-email.png", "Team Email Settings", "Per-team email configuration"),
            ("admin-sso-config.png", "SSO Configuration", "Single Sign-On settings"),
            ("admin-sso-setup.png", "SSO Setup", "SSO configuration wizard"),
            ("admin-sso-test.png", "SSO Test", "Test SSO connectivity"),
            ("admin-servicenow.png", "ServiceNow Config", "ServiceNow API settings"),
            ("admin-servicenow-setup.png", "ServiceNow Setup", "ServiceNow wizard"),
            ("admin-app-config.png", "Application Config", "Application settings"),
            ("admin-app-settings.png", "App Settings", "General application settings"),
            ("admin-smtp-setup.png", "SMTP Setup", "SMTP configuration wizard"),
        ]
    },
    "Shift Time Configuration": {
        "description": "Configure shift times for teams",
        "files": [
            ("admin-shift-time-config.png", "Shift Times", "Shift time configuration"),
            ("admin-shift-time-edit.png", "Edit Shift Time", "Modify shift times"),
        ]
    },
    "Session Monitoring": {
        "description": "Super Admin - Monitor active sessions",
        "files": [
            ("admin-active-sessions.png", "Active Sessions", "Currently logged in users"),
            ("admin-session-actions.png", "Session Actions", "Terminate user sessions"),
        ]
    },
    "Secrets Management": {
        "description": "Super Admin - Manage application secrets",
        "files": [
            ("admin-secrets-management.png", "Secrets Management", "Manage secure credentials"),
            ("admin-add-secret.png", "Add Secret", "Add new secret"),
        ]
    },
    "Incident Response Logs": {
        "description": "Track incident response performance",
        "files": [
            ("admin-incident-logs.png", "Incident Logs", "Response log history"),
            ("admin-incident-filters.png", "Incident Filters", "Filter incident logs"),
        ]
    },
    "System Health & Monitoring": {
        "description": "Monitor system health and performance",
        "files": [
            ("admin-system-health.png", "System Health", "System health dashboard"),
            ("admin-email-monitoring.png", "Email Monitoring", "Email delivery tracking"),
        ]
    },
    "Vendor Management": {
        "description": "Manage vendor contact information",
        "files": [
            ("admin-vendor-management.png", "Vendor List", "All vendor contacts"),
            ("admin-add-vendor.png", "Add Vendor", "New vendor form"),
            ("admin-vendor-upload.png", "Vendor Upload", "Bulk vendor import"),
        ]
    },
    "Escalation Matrix": {
        "description": "Manage escalation contacts",
        "files": [
            ("admin-escalation-management.png", "Escalation Matrix", "Escalation contacts"),
            ("admin-add-escalation.png", "Add Escalation", "Add escalation contact"),
            ("admin-escalation-upload.png", "Escalation Upload", "Bulk escalation import"),
        ]
    },
    "Roster Administration": {
        "description": "Manage shift rosters",
        "files": [
            ("admin-roster-team.png", "Team Roster", "Roster management view"),
            ("admin-roster-upload.png", "Roster Upload", "Upload roster page"),
            ("admin-roster-upload-interface.png", "Upload Interface", "Roster upload form"),
            ("admin-roster-review.png", "Roster Review", "Review imported roster"),
            ("admin-roster-manual.png", "Manual Entry", "Manual roster entry"),
            ("admin-shift-allowance.png", "Shift Allowance", "Generate allowance report"),
            ("admin-roster-template.png", "Roster Template", "Download template"),
        ]
    },
    "Reporting": {
        "description": "Admin reporting features",
        "files": [
            ("admin-reports-account.png", "Account Reports", "Multi-team reports view"),
            ("admin-reports-team.png", "Team Reports", "Team-specific reports"),
            ("admin-delete-draft.png", "Delete Draft", "Delete draft confirmation"),
        ]
    },
    "Audit & Compliance": {
        "description": "Audit logs and compliance monitoring",
        "files": [
            ("admin-audit-global.png", "Global Audit", "System-wide audit logs"),
            ("admin-audit-account.png", "Account Audit", "Account-level audit logs"),
            ("admin-audit-viewer.png", "Audit Viewer", "Detailed audit log view"),
        ]
    },
}


def create_document(title, sections, screenshots_dir, output_file):
    """Create a Word document with screenshots."""
    if not DOCX_AVAILABLE:
        print("[ERROR] Cannot create document - python-docx not installed")
        return False
    
    doc = Document()
    
    # Set up styles
    title_style = doc.styles['Title']
    title_style.font.size = Pt(28)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(30, 60, 114)
    
    heading1_style = doc.styles['Heading 1']
    heading1_style.font.size = Pt(18)
    heading1_style.font.bold = True
    heading1_style.font.color.rgb = RGBColor(30, 60, 114)
    
    heading2_style = doc.styles['Heading 2']
    heading2_style.font.size = Pt(14)
    heading2_style.font.bold = True
    heading2_style.font.color.rgb = RGBColor(42, 82, 152)
    
    # Title Page
    doc.add_paragraph()
    doc.add_paragraph()
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(36)
    title_run.font.color.rgb = RGBColor(30, 60, 114)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("Shift Handover Application")
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = RGBColor(128, 128, 128)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading("Table of Contents", level=1)
    toc_para = doc.add_paragraph()
    for idx, section_name in enumerate(sections.keys(), 1):
        toc_para.add_run(f"{idx}. {section_name}\n")
    
    doc.add_page_break()
    
    # Content sections
    section_num = 0
    for section_name, section_data in sections.items():
        section_num += 1
        
        # Section header
        doc.add_heading(f"{section_num}. {section_name}", level=1)
        
        # Section description
        desc_para = doc.add_paragraph()
        desc_run = desc_para.add_run(section_data["description"])
        desc_run.font.italic = True
        desc_run.font.color.rgb = RGBColor(100, 100, 100)
        
        doc.add_paragraph()
        
        # Screenshots
        for filename, screenshot_title, description in section_data["files"]:
            filepath = screenshots_dir / filename
            
            # Screenshot title
            doc.add_heading(screenshot_title, level=2)
            
            # Description
            desc = doc.add_paragraph(description)
            desc.paragraph_format.space_after = Pt(12)
            
            # Add image if exists
            if filepath.exists():
                try:
                    doc.add_picture(str(filepath), width=Inches(6))
                    last_para = doc.paragraphs[-1]
                    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    error_para = doc.add_paragraph()
                    error_run = error_para.add_run(f"[Image: {filename}]")
                    error_run.font.color.rgb = RGBColor(200, 0, 0)
                    print(f"  [WARNING] Could not add image {filename}: {e}")
            else:
                missing_para = doc.add_paragraph()
                missing_run = missing_para.add_run(f"[Screenshot not found: {filename}]")
                missing_run.font.color.rgb = RGBColor(200, 0, 0)
            
            doc.add_paragraph()
        
        doc.add_page_break()
    
    # Save document
    doc.save(output_file)
    print(f"[OK] Created: {output_file}")
    return True


def convert_to_pdf(docx_file, pdf_file):
    """Convert Word document to PDF."""
    if not PDF_AVAILABLE:
        print("[ERROR] Cannot convert to PDF - docx2pdf not installed")
        print("   Install with: pip install docx2pdf")
        print("   Note: Requires Microsoft Word on Windows or LibreOffice on Mac/Linux")
        return False
    
    try:
        convert(docx_file, pdf_file)
        print(f"[OK] Created: {pdf_file}")
        return True
    except Exception as e:
        print(f"[ERROR] PDF conversion failed: {e}")
        print("   Make sure Microsoft Word or LibreOffice is installed")
        return False


def main():
    print("=" * 60)
    print("[DOC] Shift Handover Documentation Generator")
    print("=" * 60)
    print()
    
    # Get paths
    script_dir = Path(__file__).parent
    screenshots_dir = script_dir / "screenshots"
    
    if not screenshots_dir.exists():
        print(f"[ERROR] Screenshots directory not found: {screenshots_dir}")
        return
    
    # Count available screenshots
    available_screenshots = list(screenshots_dir.glob("*.png"))
    print(f"[INFO] Found {len(available_screenshots)} screenshots in {screenshots_dir}")
    print()
    
    # Generate User Guide
    print("[USER GUIDE] Generating User Guide...")
    user_guide_docx = script_dir / "Shift_Handover_User_Guide.docx"
    user_guide_pdf = script_dir / "Shift_Handover_User_Guide.pdf"
    
    if create_document("User Guide", USER_GUIDE_SECTIONS, screenshots_dir, user_guide_docx):
        convert_to_pdf(user_guide_docx, user_guide_pdf)
    
    print()
    
    # Generate Admin Guide
    print("[ADMIN GUIDE] Generating Admin Guide...")
    admin_guide_docx = script_dir / "Shift_Handover_Admin_Guide.docx"
    admin_guide_pdf = script_dir / "Shift_Handover_Admin_Guide.pdf"
    
    if create_document("Administrator Guide", ADMIN_GUIDE_SECTIONS, screenshots_dir, admin_guide_docx):
        convert_to_pdf(admin_guide_docx, admin_guide_pdf)
    
    print()
    print("=" * 60)
    print("[SUCCESS] Documentation generation complete!")
    print()
    print("Generated files:")
    print(f"  - {user_guide_docx}")
    print(f"  - {user_guide_pdf}")
    print(f"  - {admin_guide_docx}")
    print(f"  - {admin_guide_pdf}")
    print("=" * 60)


if __name__ == "__main__":
    main()

